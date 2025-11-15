import os
from docx import Document
import openpyxl

# --- Configuración de Rutas ---
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
DOCS_DIR = os.path.join(BASE_DIR, 'Software II')

# --- Nombres de los archivos (ajusta las extensiones si es necesario) ---
# Asumimos que la rúbrica es .docx, si es .pdf u otro, este script no podrá leerlo.
RUBRICA_FILE = 'Rúbrica de Evaluación 3.docx' 
GUIA_FILE = 'Guía - Plan de Pruebas.docx'
XLSX_FILE = 'R001-Casos de Prueba.xlsx'

def print_separator(title):
    print("\n" + "="*20 + f" {title} " + "="*20 + "\n")

def read_docx(file_path, file_title):
    """Lee y muestra el contenido de un archivo .docx, incluyendo tablas."""
    print_separator(f"INICIO: {file_title}")
    try:
        if not os.path.exists(file_path):
            print(f"ERROR: El archivo no se encontró en la ruta: {file_path}")
            print("Por favor, verifica que el nombre y la extensión del archivo son correctos.")
            return

        document = Document(file_path)
        
        print("--- Contenido del Documento ---")
        for para in document.paragraphs:
            if para.text.strip(): # Evitar imprimir líneas vacías
                print(para.text)
        
        if document.tables:
            print("\n--- Contenido de las Tablas ---")
            for i, table in enumerate(document.tables):
                print(f"\n--- Tabla {i+1} ---")
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    print(' | '.join(row_data))

    except Exception as e:
        print(f"Ocurrió un error al leer el archivo {file_title}: {e}")
        print("Asegúrate de que el archivo no esté dañado y que la extensión sea .docx")
    finally:
        print_separator(f"FIN: {file_title}")

def read_xlsx_headers(file_path, file_title):
    """Lee y muestra solo la primera fila (cabeceras) de un archivo .xlsx."""
    print_separator(f"INICIO: {file_title}")
    try:
        if not os.path.exists(file_path):
            print(f"ERROR: El archivo no se encontró en la ruta: {file_path}")
            return

        workbook = openpyxl.load_workbook(file_path)
        sheet = workbook.active
        
        print("--- Cabeceras del Excel (Columnas) ---")
        header_row = [cell.value for cell in sheet[1] if cell.value]
        print(' | '.join(header_row))

    except Exception as e:
        print(f"Ocurrió un error al leer el archivo {file_title}: {e}")
    finally:
        print_separator(f"FIN: {file_title}")

if __name__ == "__main__":
    # Construir rutas completas
    rubrica_path = os.path.join(DOCS_DIR, RUBRICA_FILE)
    guia_path = os.path.join(DOCS_DIR, GUIA_FILE)
    xlsx_path = os.path.join(DOCS_DIR, XLSX_FILE)

    # Leer y mostrar los contenidos
    read_docx(rubrica_path, "Rúbrica de Evaluación 3")
    read_docx(guia_path, "Guía - Plan de Pruebas")
    read_xlsx_headers(xlsx_path, "Template de Casos de Prueba")

    print("\nProceso de extracción finalizado.")
    print("Por favor, copia todo el texto generado arriba y pégalo en el chat.")
